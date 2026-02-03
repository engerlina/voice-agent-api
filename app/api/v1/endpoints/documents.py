"""Document management endpoints for RAG knowledge base."""

import io
import re
import uuid
from datetime import datetime
from typing import Any, List, Optional
from urllib.parse import urlparse

import httpx
from fastapi import APIRouter, BackgroundTasks, Depends, File, HTTPException, UploadFile, status
from pydantic import BaseModel, HttpUrl
from sqlalchemy import delete, desc, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.v1.endpoints.auth import get_current_user
from app.core.config import settings
from app.core.database import get_db
from app.core.logging import logger
from app.models.document import Document, DocumentStatus, DocumentType
from app.models.user import User
from app.services.rag_service import RAGService

# File parsing imports
try:
    from pypdf import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Supported file extensions and their types
SUPPORTED_EXTENSIONS = {
    ".pdf": DocumentType.PDF,
    ".txt": DocumentType.TXT,
    ".md": DocumentType.TXT,
    ".docx": DocumentType.DOCX,
}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

router = APIRouter()


# ============== Schemas ==============


class DocumentCreate(BaseModel):
    """Create document request."""

    name: str
    description: Optional[str] = None
    content: str
    metadata: Optional[dict[str, Any]] = None


class DocumentResponse(BaseModel):
    """Document response schema."""

    id: str
    name: str
    description: Optional[str]
    status: str
    chunk_count: int
    created_at: datetime
    updated_at: datetime

    class Config:
        from_attributes = True


class DocumentDetailResponse(DocumentResponse):
    """Detailed document response with content preview."""

    content_preview: Optional[str]  # First 500 chars
    error_message: Optional[str]


class SearchRequest(BaseModel):
    """RAG search request."""

    query: str
    top_k: int = 5


class SearchResult(BaseModel):
    """Search result."""

    document_name: str
    content: str
    similarity: float


class SearchResponse(BaseModel):
    """Search response."""

    query: str
    results: List[SearchResult]


class UrlImportRequest(BaseModel):
    """URL import request."""

    url: HttpUrl
    name: Optional[str] = None  # Optional custom name, defaults to page title


# ============== File Parsing Helpers ==============


def parse_pdf_content(file_bytes: bytes) -> str:
    """Extract text from PDF file."""
    if not PDF_SUPPORT:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="PDF parsing not available - pypdf not installed",
        )

    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    return "\n\n".join(text_parts)


def parse_docx_content(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    if not DOCX_SUPPORT:
        raise HTTPException(
            status_code=status.HTTP_501_NOT_IMPLEMENTED,
            detail="DOCX parsing not available - python-docx not installed",
        )

    doc = DocxDocument(io.BytesIO(file_bytes))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    return "\n\n".join(text_parts)


def parse_text_content(file_bytes: bytes) -> str:
    """Extract text from TXT/MD file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Unable to decode text file - unsupported encoding",
            )


def get_file_extension(filename: str) -> str:
    """Get lowercase file extension from filename."""
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


async def scrape_url_content(url: str) -> tuple[str, str]:
    """Scrape content from URL using Browserless API.

    Returns tuple of (title, content).
    """
    if not settings.browserless_api_key:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="URL scraping not configured - Browserless API key required",
        )

    # Use Browserless /content endpoint for text extraction
    browserless_url = f"https://chrome.browserless.io/content?token={settings.browserless_api_key}"

    payload = {
        "url": url,
        "waitFor": 2000,  # Wait 2s for JS to load
    }

    async with httpx.AsyncClient(timeout=30.0) as client:
        try:
            response = await client.post(browserless_url, json=payload)
            response.raise_for_status()
            html_content = response.text
        except httpx.TimeoutException:
            raise HTTPException(
                status_code=status.HTTP_408_REQUEST_TIMEOUT,
                detail="URL scraping timed out",
            )
        except httpx.HTTPStatusError as e:
            logger.error("browserless_error", status=e.response.status_code, detail=e.response.text)
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=f"Failed to fetch URL: {e.response.status_code}",
            )
        except Exception as e:
            logger.error("url_scrape_failed", url=url, error=str(e))
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to scrape URL: {str(e)}",
            )

    # Extract title from HTML
    title_match = re.search(r"<title[^>]*>([^<]+)</title>", html_content, re.IGNORECASE)
    title = title_match.group(1).strip() if title_match else urlparse(url).netloc

    # Strip HTML tags to get plain text
    # Remove script and style elements first
    clean_html = re.sub(r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL | re.IGNORECASE)
    clean_html = re.sub(r"<style[^>]*>.*?</style>", "", clean_html, flags=re.DOTALL | re.IGNORECASE)
    # Remove all HTML tags
    text_content = re.sub(r"<[^>]+>", " ", clean_html)
    # Clean up whitespace
    text_content = re.sub(r"\s+", " ", text_content).strip()
    # Decode HTML entities
    text_content = text_content.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"')

    return title, text_content


# ============== Background Tasks ==============


async def process_document_background(
    document_id: uuid.UUID,
    content: str,
    db: AsyncSession,
):
    """Background task to process document and create embeddings."""
    try:
        rag_service = RAGService(db)
        chunk_count = await rag_service.process_document(document_id, content)

        logger.info(
            "document_processed",
            document_id=str(document_id),
            chunk_count=chunk_count,
        )
    except Exception as e:
        logger.error(
            "document_processing_failed",
            document_id=str(document_id),
            error=str(e),
        )
        # Update document status to failed
        result = await db.execute(select(Document).where(Document.id == document_id))
        doc = result.scalar_one_or_none()
        if doc:
            doc.status = DocumentStatus.FAILED
            doc.error_message = str(e)
            await db.commit()


# ============== Endpoints ==============


@router.post("", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def create_document(
    request: DocumentCreate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Create a document and queue it for RAG processing.

    The document content will be chunked and embedded in the background.
    Check the document status to see when processing is complete.
    """
    if not settings.pinecone_api_key or not settings.pinecone_url:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG not configured - Pinecone settings required",
        )

    if not request.content.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document content cannot be empty",
        )

    # Create document record
    document = Document(
        user_id=current_user.id,
        name=request.name,
        description=request.description,
        content=request.content,
        doc_type=DocumentType.MANUAL,
        status=DocumentStatus.PENDING,
        extra_data=request.metadata,
    )

    db.add(document)
    await db.commit()
    await db.refresh(document)

    logger.info(
        "document_created",
        document_id=str(document.id),
        user_id=current_user.id,
        name=request.name,
    )

    # Process in background (chunk + embed + store in Pinecone)
    # Note: We need a new db session for background task
    # For now, process synchronously to ensure it completes
    try:
        rag_service = RAGService(db)
        chunk_count = await rag_service.process_document(document.id, request.content)
        await db.refresh(document)
    except Exception as e:
        logger.error("document_processing_failed", error=str(e))
        document.status = DocumentStatus.FAILED
        document.error_message = str(e)
        await db.commit()

    return DocumentResponse(
        id=str(document.id),
        name=document.name,
        description=document.description,
        status=document.status.value,
        chunk_count=document.chunk_count,
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    description: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Upload a file (PDF, DOCX, TXT, MD) and process it for RAG.

    Supported formats:
    - PDF: Extracts text from all pages
    - DOCX: Extracts text from paragraphs
    - TXT/MD: Reads as plain text

    Max file size: 10MB
    """
    if not settings.pinecone_api_key or not settings.pinecone_url:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG not configured - Pinecone settings required",
        )

    # Validate file extension
    ext = get_file_extension(file.filename or "")
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS.keys())}",
        )

    # Read file content
    file_bytes = await file.read()

    # Validate file size
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024 * 1024)}MB",
        )

    if len(file_bytes) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File is empty",
        )

    # Parse content based on file type
    doc_type = SUPPORTED_EXTENSIONS[ext]
    try:
        if doc_type == DocumentType.PDF:
            content = parse_pdf_content(file_bytes)
        elif doc_type == DocumentType.DOCX:
            content = parse_docx_content(file_bytes)
        else:  # TXT or MD
            content = parse_text_content(file_bytes)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("file_parsing_failed", filename=file.filename, error=str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse file: {str(e)}",
        )

    if not content.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No text content could be extracted from the file",
        )

    # Create document record
    document = Document(
        user_id=current_user.id,
        name=file.filename or "Uploaded Document",
        description=description,
        content=content,
        doc_type=doc_type,
        status=DocumentStatus.PENDING,
    )

    db.add(document)
    await db.commit()
    await db.refresh(document)

    logger.info(
        "document_uploaded",
        document_id=str(document.id),
        user_id=current_user.id,
        filename=file.filename,
        doc_type=doc_type.value,
        content_length=len(content),
    )

    # Process through RAG pipeline
    try:
        rag_service = RAGService(db)
        chunk_count = await rag_service.process_document(document.id, content)
        await db.refresh(document)
    except Exception as e:
        logger.error("document_processing_failed", error=str(e))
        document.status = DocumentStatus.FAILED
        document.error_message = str(e)
        await db.commit()

    return DocumentResponse(
        id=str(document.id),
        name=document.name,
        description=document.description,
        status=document.status.value,
        chunk_count=document.chunk_count,
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


@router.post("/url", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def import_url(
    request: UrlImportRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Import content from a URL and process it for RAG.

    Uses Browserless to fetch and extract text content from web pages.
    Supports JavaScript-rendered content.
    """
    if not settings.pinecone_api_key or not settings.pinecone_url:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG not configured - Pinecone settings required",
        )

    url_str = str(request.url)
    logger.info("url_import_started", url=url_str, user_id=current_user.id)

    # Scrape the URL
    title, content = await scrape_url_content(url_str)

    if not content.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No text content could be extracted from the URL",
        )

    # Use custom name or scraped title
    doc_name = request.name or title or urlparse(url_str).netloc

    # Create document record
    document = Document(
        user_id=current_user.id,
        name=doc_name,
        description=f"Imported from: {url_str}",
        content=content,
        doc_type=DocumentType.URL,
        status=DocumentStatus.PENDING,
        extra_data={"source_url": url_str},
    )

    db.add(document)
    await db.commit()
    await db.refresh(document)

    logger.info(
        "url_imported",
        document_id=str(document.id),
        user_id=current_user.id,
        url=url_str,
        content_length=len(content),
    )

    # Process through RAG pipeline
    try:
        rag_service = RAGService(db)
        chunk_count = await rag_service.process_document(document.id, content)
        await db.refresh(document)
    except Exception as e:
        logger.error("document_processing_failed", error=str(e))
        document.status = DocumentStatus.FAILED
        document.error_message = str(e)
        await db.commit()

    return DocumentResponse(
        id=str(document.id),
        name=document.name,
        description=document.description,
        status=document.status.value,
        chunk_count=document.chunk_count,
        created_at=document.created_at,
        updated_at=document.updated_at,
    )


@router.get("", response_model=List[DocumentResponse])
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    limit: int = 50,
    offset: int = 0,
):
    """List all documents for the current user."""
    result = await db.execute(
        select(Document)
        .where(Document.user_id == current_user.id)
        .order_by(desc(Document.created_at))
        .limit(limit)
        .offset(offset)
    )
    documents = result.scalars().all()

    return [
        DocumentResponse(
            id=str(doc.id),
            name=doc.name,
            description=doc.description,
            status=doc.status.value,
            chunk_count=doc.chunk_count,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
        )
        for doc in documents
    ]


@router.get("/{document_id}", response_model=DocumentDetailResponse)
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Get a specific document by ID."""
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format",
        )

    result = await db.execute(
        select(Document)
        .where(Document.id == doc_uuid)
        .where(Document.user_id == current_user.id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    # Content preview (first 500 chars)
    content_preview = None
    if document.content:
        content_preview = document.content[:500]
        if len(document.content) > 500:
            content_preview += "..."

    return DocumentDetailResponse(
        id=str(document.id),
        name=document.name,
        description=document.description,
        status=document.status.value,
        chunk_count=document.chunk_count,
        created_at=document.created_at,
        updated_at=document.updated_at,
        content_preview=content_preview,
        error_message=document.error_message,
    )


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Delete a document and its vectors from Pinecone."""
    try:
        doc_uuid = uuid.UUID(document_id)
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid document ID format",
        )

    result = await db.execute(
        select(Document)
        .where(Document.id == doc_uuid)
        .where(Document.user_id == current_user.id)
    )
    document = result.scalar_one_or_none()

    if not document:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found",
        )

    # Delete vectors from Pinecone
    try:
        rag_service = RAGService(db)
        await rag_service.delete_document(doc_uuid, current_user.id)
    except Exception as e:
        logger.warning(
            "pinecone_delete_failed",
            document_id=document_id,
            error=str(e),
        )
        # Continue with database deletion even if Pinecone fails

    # Delete from database
    await db.execute(delete(Document).where(Document.id == doc_uuid))
    await db.commit()

    logger.info(
        "document_deleted",
        document_id=document_id,
        user_id=current_user.id,
    )

    return None


@router.post("/search", response_model=SearchResponse)
async def search_documents(
    request: SearchRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Search documents using RAG vector similarity."""
    if not settings.pinecone_api_key or not settings.pinecone_url:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="RAG not configured - Pinecone settings required",
        )

    if not request.query.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Search query cannot be empty",
        )

    logger.info(
        "rag_search",
        user_id=current_user.id,
        query=request.query[:50],
        top_k=request.top_k,
    )

    try:
        rag_service = RAGService(db)
        results = await rag_service.search(
            user_id=current_user.id,
            query=request.query,
            top_k=request.top_k,
            similarity_threshold=0.5,  # Lower threshold for more results
        )

        return SearchResponse(
            query=request.query,
            results=[
                SearchResult(
                    document_name=r.get("document_name", "Unknown"),
                    content=r.get("content", ""),
                    similarity=round(r.get("similarity", 0) * 100, 1),  # Convert to percentage
                )
                for r in results
            ],
        )
    except Exception as e:
        logger.error("rag_search_failed", error=str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Search failed: {str(e)}",
        )
